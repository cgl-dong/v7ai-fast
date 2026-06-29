"""Document indexing service — splits, embeds, and stores in pgvector.

Powered by LlamaIndex text splitters + Chinese-aware custom splitters:
  - recursive  — 多级分隔符递归切分 (新默认, 通用首选)
  - sentence   — 句子级智能切分 (PDF/DOCX)
  - paragraph  — 按段落切分, 保留篇章结构 (MD/TXT)
  - token      — Token 级切分, 对齐模型上下文窗口
  - fixed      — 固定长度切分 (CSV 等)
  - excel      — Excel 行分组切分
  - section    — 面向法规文档, 按章/节/条切分
  - qa         — 面向问答对, 按问/答边界分组
  - semantic   — 基于 embedding 相似度的语义切分

Embedding model: BAAI/bge-base-zh-v1.5 (768 dims)."""
import io
import json
import logging
from typing import List, Iterable, Optional
from sqlalchemy.orm import Session

from app.core.database import KnowledgeFile, DocumentChunk
from app.core.settings import settings
from app.services.knowledge import KnowledgeService
from app.services.embedding import embed_texts, embed_query
from app.services.chunking import split_text, CHUNK_CONFIG
from app.services.hybrid_search import tokenize, tokens_to_str

# Lazy import to avoid circular deps — skills are loaded on demand
_skill_registry = None


def _get_registry():
    global _skill_registry
    if _skill_registry is None:
        from app.services.skill_base import registry
        # Trigger skill auto-discovery
        import app.services.skills  # noqa: F401
        _skill_registry = registry
    return _skill_registry

logger = logging.getLogger(__name__)


class Indexer:
    """Handles document loading, splitting, embedding, and pgvector storage."""

    def __init__(self, db: Session):
        self.db = db
        self.knowledge_svc = KnowledgeService(db)

    def index_file(self, file_id: int, strategy: str = None, chunk_size: int = None,
                   chunk_overlap: int = None, skills: Optional[List[str]] = None) -> dict:
        """Index a single file: (skills) → load → parse → split → embed → store.

        Optionally override chunking strategy/size/overlap per call.
        If skills are specified, the skill pipeline transforms the file content
        before parsing (original MinIO file is unchanged).
        """
        record = self.db.query(KnowledgeFile).filter(KnowledgeFile.id == file_id).first()
        if not record:
            return {"error": "文件不存在"}

        applied_skills = []
        try:
            record.status = "indexing"
            record.error_msg = None
            self.db.commit()

            # 1. Load content from MinIO
            import time
            t0 = time.time()
            content_bytes, _, _ = self.knowledge_svc.get_file_content(file_id)

            # 1.5 Apply skill pipeline if specified
            parse_filename = record.filename
            parse_file_type = record.file_type
            converted_file_id = None
            if skills:
                registry = _get_registry()
                try:
                    converted_bytes, converted_name, converted_type = registry.run_pipeline(
                        content=content_bytes,
                        filename=record.filename,
                        file_type=record.file_type,
                        skill_names=skills,
                    )
                    applied_skills = skills
                    logger.info(
                        f"[index] skills applied: {skills}, "
                        f"result_type={converted_type}, size={len(converted_bytes)}B"
                    )

                    # Save converted file to MinIO and create a new KnowledgeFile record
                    converted_file_id = self._save_converted_file(
                        original=record,
                        content=converted_bytes,
                        filename=converted_name,
                        file_type=converted_type,
                        skills=skills,
                    )
                    # Index the converted file content
                    content_bytes = converted_bytes
                    parse_filename = converted_name
                    parse_file_type = converted_type
                    logger.info(f"[index] converted file saved: id={converted_file_id}")
                except (ValueError, RuntimeError) as e:
                    record.status = "error"
                    record.error_msg = f"技能管线失败: {e}"
                    self.db.commit()
                    return {"error": str(e)}

            text = self._parse_content(content_bytes, parse_filename, parse_file_type)
            parse_time = time.time() - t0
            logger.info(f"[index] file={record.filename} (id={file_id}, type={record.file_type}, "
                        f"size={record.file_size}B), parsed {len(text)} chars in {parse_time:.1f}s")

            # 1.5 Validate content quality — reject binary garbage
            err = self._validate_content(text, record.filename)
            if err:
                record.status = "error"
                record.error_msg = err
                self.db.commit()
                return {"error": err}

            # 2. Get chunk config for this file type (with optional overrides)
            from app.services.chunking import ChunkConfig
            cfg = CHUNK_CONFIG.get(record.file_type, CHUNK_CONFIG["default"])
            effective_config = ChunkConfig(
                strategy=strategy or cfg.strategy,
                chunk_size=chunk_size or cfg.chunk_size,
                chunk_overlap=chunk_overlap or cfg.chunk_overlap,
            )

            # 3. Split into chunks using LlamaIndex strategy
            chunks = split_text(text, record.file_type, config=effective_config)

            if not chunks:
                record.status = "error"
                record.error_msg = "文档内容为空，无法分片"
                self.db.commit()
                return {"error": "文档内容为空"}

            # 4. Delete old chunks
            self.db.query(DocumentChunk).filter(DocumentChunk.file_id == file_id).delete()

            # 5. Generate embeddings and store
            texts = [c["content"] for c in chunks]
            logger.info(f"Generating embeddings for {len(texts)} chunks "
                        f"(strategy={effective_config.strategy}, size={effective_config.chunk_size}, "
                        f"overlap={effective_config.chunk_overlap})...")
            embeddings = embed_texts(texts)

            for i, chunk in enumerate(chunks):
                metadata = chunk.get("metadata", {})
                metadata["chunk_strategy"] = effective_config.strategy
                metadata["chunk_size"] = effective_config.chunk_size
                metadata["chunk_overlap"] = effective_config.chunk_overlap
                if applied_skills:
                    metadata["skills_applied"] = applied_skills
                # Tokenize for BM25 keyword search
                tokens_str = tokens_to_str(tokenize(chunk["content"]))
                dc = DocumentChunk(
                    file_id=file_id,
                    chunk_index=i,
                    content=chunk["content"],
                    tokens=tokens_str,
                    embedding=embeddings[i],
                    metadata_json=json.dumps(metadata, ensure_ascii=False),
                )
                self.db.add(dc)

            # 6. Update status
            record.status = "indexed"
            record.chunk_count = len(chunks)
            record.error_msg = None
            self.db.commit()

            logger.info(f"Indexed file {file_id}: {len(chunks)} chunks (strategy={effective_config.strategy})")
            result = {"success": True, "chunks": len(chunks), "strategy": effective_config.strategy}
            if applied_skills:
                result["skills_applied"] = applied_skills
            if converted_file_id:
                result["converted_file_id"] = converted_file_id
            return result

        except Exception as e:
            logger.error(f"Failed to index file {file_id}: {e}")
            record.status = "error"
            record.error_msg = str(e)
            self.db.commit()
            raise

    def _save_converted_file(
        self, original: KnowledgeFile, content: bytes,
        filename: str, file_type: str, skills: List[str],
    ) -> int:
        """Save skill pipeline output to MinIO and create a KnowledgeFile record.

        Returns the new file's ID.
        """
        import uuid
        from app.core.settings import settings
        from minio import Minio

        ext = file_type
        stored_name = f"{uuid.uuid4().hex}.{ext}"
        file_size = len(content)

        # Upload to MinIO
        client = Minio(
            endpoint=settings.minio_endpoint,
            access_key=settings.minio_access_key,
            secret_key=settings.minio_secret_key,
            secure=settings.minio_secure,
        )
        if not client.bucket_exists(settings.minio_bucket):
            client.make_bucket(settings.minio_bucket)
        client.put_object(
            bucket_name=settings.minio_bucket,
            object_name=stored_name,
            data=io.BytesIO(content),
            length=file_size,
            content_type="application/octet-stream",
        )

        # Create DB record linking to original
        new_record = KnowledgeFile(
            filename=filename,
            stored_name=stored_name,
            file_type=file_type,
            file_size=file_size,
            file_path=stored_name,
            status="uploaded",
            uploader=original.uploader or "skill-pipeline",
            kb_id=original.kb_id,
            user_id=original.user_id,
            chunk_strategy=original.chunk_strategy,
            chunk_size=original.chunk_size,
            chunk_overlap=original.chunk_overlap,
            source_file_id=original.id,
            skills_applied=json.dumps(skills, ensure_ascii=False),
        )
        self.db.add(new_record)
        self.db.commit()
        self.db.refresh(new_record)
        return new_record.id

    def _parse_content(self, content: bytes, filename: str, file_type: str) -> str:
        """Parse file bytes into plain text."""
        ext = file_type
        logger.debug(f"[parse] {filename}: type={ext}, bytes={len(content)}")

        text = ""
        if ext in ("txt", "md", "csv"):
            text = content.decode("utf-8", errors="replace")
            logger.debug(f"[parse] {filename}: decoded as text, {len(text)} chars")

        elif ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            text = "\n\n".join(pages)
            logger.debug(f"[parse] {filename}: PDF parsed, {len(reader.pages)} pages, {len(text)} chars")

        elif ext == "xlsx":
            try:
                from openpyxl import load_workbook
                wb = load_workbook(io.BytesIO(content), read_only=True)
                rows = []
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    rows.append(f"[Sheet: {sheet_name}]")
                    for row in sheet.iter_rows(values_only=True):
                        row_str = "\t".join(str(c) if c is not None else "" for c in row)
                        if row_str.strip():
                            rows.append(row_str)
                wb.close()
                text = "\n".join(rows)
                logger.debug(f"[parse] {filename}: XLSX parsed, {len(wb.sheetnames)} sheets, {len(rows)} rows, {len(text)} chars")
            except Exception as e:
                logger.warning(f"[parse] {filename}: XLSX parse failed: {e}, falling back to text decode")
                text = content.decode("utf-8", errors="replace")

        elif ext == "docx":
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                logger.debug(f"[parse] {filename}: DOCX parsed, {len(doc.paragraphs)} paragraphs, {len(text)} chars")
            except ImportError:
                logger.warning(f"[parse] {filename}: python-docx not installed")
                text = ""
            except Exception as e:
                logger.warning(f"[parse] {filename}: DOCX parse failed: {e}, falling back to text decode")
                text = content.decode("utf-8", errors="replace")

        else:
            logger.warning(f"[parse] {filename}: unknown type '{ext}', raw decode")
            text = content.decode("utf-8", errors="replace")

        # Strip NUL and non-printable control chars that break PostgreSQL
        text = text.replace("\x00", "").replace("\r", "\n")
        return text

    def _validate_content(self, text: str, filename: str) -> str | None:
        """Validate that parsed content looks like readable text, not binary garbage.

        Returns error message string if content is invalid, None if OK.
        """
        if not text or not text.strip():
            logger.warning(f"[validate] {filename}: empty content")
            return "文档解析后内容为空"

        total = len(text)
        printable = sum(
            1 for c in text
            if c.isprintable() or c in "\n\t\r" or
               ('\u4e00' <= c <= '\u9fff') or
               ('\u3000' <= c <= '\u303f') or
               ('\uff00' <= c <= '\uffef')
        )
        ratio = printable / total if total > 0 else 0
        logger.debug(f"[validate] {filename}: printable_ratio={ratio:.1%}, total_chars={total}")

        if ratio < 0.5:
            logger.warning(f"[validate] {filename}: low printable ratio {ratio:.1%}, rejecting")
            return f"文件内容可读率仅 {ratio:.1%}，可能是二进制文件或加密文件"

        return None

    def search_chunks(self, query: str, top_k: int = 5, kb_id: int = None, user_id: int = None) -> List[dict]:
        """Search chunks with hybrid retrieval + Cross-Encoder rerank.

        Pipeline: Dense (pgvector) + BM25 (PG tsvector) → RRF fusion → rerank → top_k
        """
        mode = (settings.rag_search_mode or "hybrid").lower()
        if mode == "dense":
            rows = self._dense_search(query, top_k, kb_id=kb_id, user_id=user_id)
        else:
            dense_k = max(top_k, settings.rag_dense_candidates)
            bm25_k = max(top_k, settings.rag_bm25_candidates)
            dense_rows = self._dense_search(query, dense_k, kb_id=kb_id, user_id=user_id)
            try:
                bm25_rows = self._bm25_search(query, bm25_k, kb_id=kb_id, user_id=user_id)
            except Exception as e:
                logger.warning(f"[search] BM25 failed, dense only: {e}")
                bm25_rows = []
            rows = _rrf_merge([dense_rows, bm25_rows], top_k=max(top_k, 20), rrf_k=settings.rag_rrf_k)

        # Cross-Encoder rerank
        if settings.rag_rerank_enabled and rows:
            from app.services.rerank import rerank
            rows = rerank(query, rows, top_n=top_k)
        else:
            rows = rows[:top_k]

        if rows:
            sims = [r.get("similarity", 0) for r in rows]
            logger.info(f"[search] {len(rows)} results, sims=[{', '.join(f'{s:.3f}' for s in sims)}], "
                        f"top_sim={sims[0]:.3f}, files={list(set(r['filename'] for r in rows))}")
        else:
            logger.warning(f"[search] no results for query: {query[:80]}")
        return rows

    def _base_filters(self, kb_id: int = None, user_id: int = None) -> tuple[str, dict]:
        """Build SQL filters shared by all retrieval paths."""
        clauses = []
        params = {}
        if kb_id is not None:
            clauses.append("kf.kb_id = :kb_id")
            params["kb_id"] = kb_id
        if user_id is not None:
            clauses.append("(kf.user_id = :user_id OR kf.user_id IS NULL)")
            params["user_id"] = user_id
        where_sql = f"WHERE {' AND '.join(clauses)}" if clauses else ""
        return where_sql, params

    def _row_to_dict(self, r, score_key: str = "similarity") -> dict:
        return {
            "id": r.id,
            "content": r.content,
            "filename": r.filename,
            "file_type": r.file_type,
            "chunk_index": r.chunk_index,
            "similarity": float(getattr(r, score_key, 0) or 0),
            "metadata": json.loads(r.metadata_json) if r.metadata_json else {},
        }

    def _dense_search(self, query: str, top_k: int = 5, kb_id: int = None, user_id: int = None) -> List[dict]:
        """Search document chunks by semantic similarity."""
        from sqlalchemy import text

        logger.info(f"[search] query: {query[:80]}..., top_k={top_k}, kb_id={kb_id}, user_id={user_id}")
        query_embedding = embed_query(query)
        embedding_str = "[" + ",".join(str(v) for v in query_embedding) + "]"

        where_sql, params = self._base_filters(kb_id=kb_id, user_id=user_id)
        sql = text(f"""
                SELECT dc.id, dc.content, dc.metadata_json, dc.chunk_index,
                       kf.filename, kf.file_type,
                       1 - (dc.embedding <=> :embedding) AS similarity
                FROM document_chunks dc
                JOIN knowledge_files kf ON dc.file_id = kf.id
                {where_sql}
                ORDER BY dc.embedding <=> :embedding
                LIMIT :limit
            """)
        result = self.db.execute(sql, {"embedding": embedding_str, "limit": top_k, **params})
        rows = [self._row_to_dict(r) for r in result]
        for row in rows:
            row["retrieval_source"] = "dense"
        return rows

    def _bm25_search(self, query: str, top_k: int = 5, kb_id: int = None, user_id: int = None) -> List[dict]:
        """Keyword retrieval using PostgreSQL full-text search with jieba tokenization.

        Uses pre-tokenized 'tokens' column for Chinese text, falls back to
        simple text vector on content for backward compat.
        """
        from sqlalchemy import text

        # Tokenize query with jieba (same as indexing)
        query_tokens = tokenize(query)
        query_ts = " & ".join(query_tokens) if query_tokens else query

        where_sql, params = self._base_filters(kb_id=kb_id, user_id=user_id)
        filter_prefix = f"{where_sql} AND" if where_sql else "WHERE"

        # Use tokens column if available (has jieba segmentation), else fallback to content
        sql = text(f"""
            SELECT dc.id, dc.content, dc.metadata_json, dc.chunk_index,
                   kf.filename, kf.file_type,
                   COALESCE(
                       ts_rank_cd(to_tsvector('simple', dc.tokens), to_tsquery('simple', :query)),
                       ts_rank_cd(to_tsvector('simple', dc.content), to_tsquery('simple', :query)),
                       0
                   ) AS bm25_score
            FROM document_chunks dc
            JOIN knowledge_files kf ON dc.file_id = kf.id
            {filter_prefix} (
                to_tsquery('simple', :query) @@ to_tsvector('simple', dc.tokens)
                OR to_tsquery('simple', :query) @@ to_tsvector('simple', dc.content)
            )
            ORDER BY bm25_score DESC
            LIMIT :limit
        """)
        result = self.db.execute(sql, {"query": query_ts, "limit": top_k, **params})
        rows = [self._row_to_dict(r, score_key="bm25_score") for r in result]
        for row in rows:
            row["retrieval_source"] = "bm25"
        return rows

    def _hybrid_search(self, query: str, top_k: int = 5, kb_id: int = None, user_id: int = None) -> List[dict]:
        dense_k = max(top_k, settings.rag_dense_candidates)
        bm25_k = max(top_k, settings.rag_bm25_candidates)
        dense_rows = self._dense_search(query, dense_k, kb_id=kb_id, user_id=user_id)
        try:
            bm25_rows = self._bm25_search(query, bm25_k, kb_id=kb_id, user_id=user_id)
        except Exception as e:
            logger.warning(f"[search] BM25 retrieval failed, falling back to dense only: {e}")
            bm25_rows = []
        return _rrf_merge([dense_rows, bm25_rows], top_k=top_k, rrf_k=settings.rag_rrf_k)


def _rrf_merge(result_sets: Iterable[List[dict]], top_k: int, rrf_k: int = 60) -> List[dict]:
    """Merge ranked result sets with Reciprocal Rank Fusion."""
    merged: dict[int, dict] = {}
    for results in result_sets:
        for rank, item in enumerate(results, start=1):
            chunk_id = item["id"]
            if chunk_id not in merged:
                merged[chunk_id] = {**item, "retrieval_score": 0.0, "retrieval_sources": []}
            merged[chunk_id]["retrieval_score"] += 1.0 / (rrf_k + rank)
            merged[chunk_id]["retrieval_sources"].append(item.get("retrieval_source", "retrieval"))
            merged[chunk_id]["similarity"] = max(merged[chunk_id].get("similarity", 0.0), item.get("similarity", 0.0))

    ranked = sorted(merged.values(), key=lambda x: x["retrieval_score"], reverse=True)
    return ranked[:top_k]
